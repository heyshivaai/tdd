"""
Report export: generates comprehensive DOCX Technology Due Diligence reports.

Modeled on the Horizon TDD Report format — each finding includes:
  - Severity-tagged title
  - Narrative description
  - Signal evidence chains (Signal → Document quote → Business impact)
  - "Ask the target" questions per finding
  - Blind spots with ⚠ markers
  - Pillar grades table in executive summary
  - Full signal inventory
  - Scan metadata appendix

Data sources:
  - vdr_intelligence_brief.json (VDR scan signals, batch results, rating)
  - domain_findings.json (agent deep diligence: domains, findings, chase list)
  - _scan_registry.json (scan timing and metadata)
"""
import io
import json
import logging
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

OUTPUTS_DIR = Path(__file__).parent.parent / "outputs"
DATA_DIR = Path(__file__).parent.parent / "data"

# Pillar display labels
PILLAR_LABELS = {
    "TechnologyArchitecture": "Technology & Architecture",
    "SecurityCompliance": "Security & Compliance",
    "OrganizationTalent": "Organization & Talent",
    "DataAIReadiness": "Data & AI Readiness",
    "RDSpendAssessment": "R&D Spend Assessment",
    "InfrastructureDeployment": "Infrastructure & Deployment",
    "SDLCProductManagement": "SDLC & Product Management",
}


def generate_report(company_name: str) -> Optional[io.BytesIO]:
    """
    Generate a comprehensive DOCX Technology Due Diligence report.

    Produces a report matching the Horizon TDD format with full evidence
    traceability: findings → signals → document quotes → business impact.

    Works with VDR scan data alone (Phase 0) or with both VDR + Agent
    deep diligence data (Phase 0 + 1).

    Args:
        company_name: Company name (subfolder in outputs/).

    Returns:
        BytesIO buffer containing the DOCX file, or None if generation failed.
    """
    company_dir = OUTPUTS_DIR / company_name
    if not company_dir.exists():
        logger.error("No output directory found for %s", company_name)
        return None

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        logger.error("python-docx not installed — cannot generate DOCX report")
        return None

    # ── Load all data sources ──────────────────────────────────────────────
    brief = _load_json(company_dir / "vdr_intelligence_brief.json")
    domain_data = _load_json(company_dir / "domain_findings.json")
    scan_registry = _load_json(OUTPUTS_DIR / "_scan_registry.json")
    scan_meta = scan_registry.get(company_name, {}) if scan_registry else {}

    if not brief and not domain_data:
        logger.error("No scan data found for %s", company_name)
        return None

    # Extract signals from brief
    all_signals = brief.get("signals", [])
    if not all_signals:
        # Fallback: extract from batch_results
        for batch in brief.get("batch_results", []):
            all_signals.extend(batch.get("signals", []))
    if not all_signals:
        # Fallback: extract from domain_slices
        for slice_data in brief.get("domain_slices", {}).values():
            all_signals.extend(slice_data.get("signals", []))

    # Determine deal metadata
    deal_id = brief.get("deal_id", scan_meta.get("deal_id", ""))
    sector = scan_meta.get("sector", brief.get("sector", ""))
    deal_type = scan_meta.get("deal_type", brief.get("deal_type", ""))
    overall_rating = brief.get("overall_signal_rating", "UNKNOWN")
    scan_date = brief.get("vdr_scan_timestamp", scan_meta.get("started_at", ""))[:10]
    domains = domain_data.get("domains", {}) if domain_data else {}

    # ── Build document ─────────────────────────────────────────────────────
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # ── Cover / Title ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TECHNOLOGY DUE DILIGENCE")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(15, 23, 42)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("VDR Scan Report")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(100, 116, 139)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"{deal_id}  ·  {sector}")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(71, 85, 105)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f"Scan Date: {scan_date}")
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(148, 163, 184)

    doc.add_page_break()

    # ── Executive Summary ──────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)

    if brief.get("executive_summary"):
        doc.add_paragraph(brief["executive_summary"])
    elif domains:
        # Build summary from domain data
        domain_grades = {pid: d.get("grade", "UNKNOWN") for pid, d in domains.items()}
        red_count = sum(1 for g in domain_grades.values() if g == "RED")
        yellow_count = sum(1 for g in domain_grades.values() if g == "YELLOW")
        green_count = sum(1 for g in domain_grades.values() if g == "GREEN")
        doc.add_paragraph(
            f"This report covers the technology due diligence analysis of {company_name}. "
            f"The analysis was conducted across {len(domains)} domains with "
            f"{len(all_signals)} signals extracted. "
            f"Overall rating: {overall_rating}. "
            f"Domain grades: {red_count} RED, {yellow_count} YELLOW, {green_count} GREEN."
        )
    else:
        # VDR-only summary
        pillar_counts = Counter()
        rating_counts = Counter()
        for sig in all_signals:
            pid = _get_pillar_id(sig)
            pillar_counts[pid] += 1
            rating_counts[sig.get("rating", "UNKNOWN").upper()] += 1
        doc.add_paragraph(
            f"This report covers the VDR scan analysis of {company_name}. "
            f"{len(all_signals)} signals were extracted across {len(pillar_counts)} pillars. "
            f"Signal ratings: {rating_counts.get('RED', 0)} RED, "
            f"{rating_counts.get('YELLOW', 0)} YELLOW, "
            f"{rating_counts.get('GREEN', 0)} GREEN. "
            f"Overall rating: {overall_rating}."
        )

    # ── Pillar Grades Table ────────────────────────────────────────────────
    doc.add_heading("Pillar Grades", level=2)

    if domains:
        # Agent-driven pillar grades
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Pillar", "Grade", "Findings", "Signals", "Confidence"]):
            hdr[i].text = h
            _bold_cell(hdr[i])

        for pid, dinfo in domains.items():
            plabel = dinfo.get("pillar_label", PILLAR_LABELS.get(pid, pid))
            grade = dinfo.get("grade", "UNKNOWN")
            findings_n = len(dinfo.get("findings", []))
            signals_n = len([s for s in all_signals if _get_pillar_id(s) == pid])
            confidence = dinfo.get("confidence", "—")
            row = table.add_row().cells
            row[0].text = plabel
            row[1].text = grade
            row[2].text = str(findings_n)
            row[3].text = str(signals_n)
            row[4].text = f"{confidence}%" if isinstance(confidence, (int, float)) else str(confidence)
    else:
        # VDR-only pillar signal table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Pillar", "Signals", "RED", "YELLOW"]):
            hdr[i].text = h
            _bold_cell(hdr[i])

        pillar_signals: dict[str, list] = {}
        for sig in all_signals:
            pid = _get_pillar_id(sig)
            pillar_signals.setdefault(pid, []).append(sig)

        for pid in sorted(pillar_signals.keys(), key=lambda p: -len(pillar_signals[p])):
            sigs = pillar_signals[pid]
            plabel = PILLAR_LABELS.get(pid, pid)
            red_n = sum(1 for s in sigs if s.get("rating", "").upper() == "RED")
            yellow_n = sum(1 for s in sigs if s.get("rating", "").upper() == "YELLOW")
            row = table.add_row().cells
            row[0].text = plabel
            row[1].text = str(len(sigs))
            row[2].text = str(red_n)
            row[3].text = str(yellow_n)

    # ══════════════════════════════════════════════════════════════════════
    # DOMAIN ANALYSIS — Full findings with evidence chains
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Domain Analysis", level=1)

    if domains:
        # Agent deep diligence format — full findings with evidence
        for pid, dinfo in sorted(
            domains.items(),
            key=lambda x: {"RED": 0, "YELLOW": 1, "GREEN": 2}.get(x[1].get("grade", ""), 3),
        ):
            _write_domain_section(doc, pid, dinfo, all_signals)
    elif all_signals:
        # VDR-only format — group signals by pillar, show evidence
        pillar_signals: dict[str, list] = {}
        for sig in all_signals:
            pid = _get_pillar_id(sig)
            pillar_signals.setdefault(pid, []).append(sig)

        for pid in sorted(
            pillar_signals.keys(),
            key=lambda p: -sum(1 for s in pillar_signals[p] if s.get("rating", "").upper() == "RED"),
        ):
            sigs = pillar_signals[pid]
            plabel = PILLAR_LABELS.get(pid, pid)
            red_n = sum(1 for s in sigs if s.get("rating", "").upper() == "RED")
            yellow_n = sum(1 for s in sigs if s.get("rating", "").upper() == "YELLOW")

            doc.add_heading(f"{plabel}", level=2)
            doc.add_paragraph(
                f"Signals: {len(sigs)}  |  RED: {red_n}  |  YELLOW: {yellow_n}"
            )

            # Sort signals: RED first, then YELLOW, then GREEN
            sorted_sigs = sorted(
                sigs,
                key=lambda s: {"RED": 0, "YELLOW": 1, "GREEN": 2}.get(
                    s.get("rating", "").upper(), 3
                ),
            )

            for sig in sorted_sigs:
                _write_signal_detail(doc, sig)

    # ══════════════════════════════════════════════════════════════════════
    # CHASE LIST — Questions for the target
    # ══════════════════════════════════════════════════════════════════════
    chase_list = domain_data.get("chase_list", []) if domain_data else []

    # Also build chase questions from signals if no agent chase list
    if not chase_list and all_signals:
        chase_list = _build_chase_from_signals(all_signals)

    if chase_list:
        doc.add_heading("Questions for the Target", level=1)

        # Group by pillar
        by_pillar: dict[str, list] = {}
        for q in chase_list:
            pillar = q.get("pillar_label", q.get("pillar_id", "General")) if isinstance(q, dict) else "General"
            by_pillar.setdefault(pillar, []).append(q)

        for pillar, questions in by_pillar.items():
            doc.add_heading(f"{pillar} ({len(questions)} questions)", level=3)
            for i, q in enumerate(questions, 1):
                question_text = q.get("question", str(q)) if isinstance(q, dict) else str(q)
                priority = q.get("priority", "") if isinstance(q, dict) else ""
                p = doc.add_paragraph()
                if priority:
                    run = p.add_run(f"[{priority.upper()}] ")
                    run.bold = True
                p.add_run(f"{i}. {question_text}")

    # ══════════════════════════════════════════════════════════════════════
    # SIGNAL INVENTORY — All signals in a table
    # ══════════════════════════════════════════════════════════════════════
    if all_signals:
        doc.add_heading("Signal Inventory", level=1)

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Signal ID", "Pillar", "Rating", "Title", "Source"]):
            hdr[i].text = h
            _bold_cell(hdr[i])

        for sig in sorted(all_signals, key=lambda s: (
            {"RED": 0, "YELLOW": 1, "GREEN": 2}.get(s.get("rating", "").upper(), 3),
            _get_pillar_id(s),
        )):
            row = table.add_row().cells
            row[0].text = sig.get("signal_id", sig.get("catalog_signal_id", "—"))
            row[1].text = PILLAR_LABELS.get(_get_pillar_id(sig), _get_pillar_id(sig))
            row[2].text = sig.get("rating", "—").upper()
            row[3].text = sig.get("title", sig.get("observation", ""))[:80]
            row[4].text = sig.get("source_doc", "—")

    # ══════════════════════════════════════════════════════════════════════
    # APPENDIX — Scan Metadata
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Appendix — Scan Metadata", level=1)

    doc.add_heading("Scan History", level=2)
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.style = "Light Grid Accent 1"
    hdr = meta_table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    _bold_cell(hdr[0])
    _bold_cell(hdr[1])

    meta_items = [
        ("Company", company_name),
        ("Deal ID", deal_id),
        ("Sector", sector),
        ("Deal Type", deal_type),
        ("Overall Rating", overall_rating),
        ("Total Signals", str(len(all_signals))),
        ("Scan Date", scan_date),
        ("Scan Mode", scan_meta.get("scan_mode", "full")),
        ("Documents Scanned", str(scan_meta.get("progress", {}).get("doc_count", "—"))),
        ("Batches Processed", str(scan_meta.get("progress", {}).get("batches_done", "—"))),
        ("Generated", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
    ]
    for field, value in meta_items:
        row = meta_table.add_row().cells
        row[0].text = field
        row[1].text = str(value)

    # ── Save to buffer ─────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info("Generated DOCX report for %s (%d bytes)", company_name, buf.getbuffer().nbytes)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════


def _load_json(path: Path) -> dict:
    """Load a JSON file, returning empty dict on any error."""
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _bold_cell(cell):
    """Make all runs in a table cell bold."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def _get_pillar_id(sig: dict) -> str:
    """Extract pillar ID from a signal dict (handles multiple field names)."""
    return sig.get("pillar_id") or sig.get("lens_id") or sig.get("lens") or "Unknown"


def _write_domain_section(doc, pid: str, dinfo: dict, all_signals: list):
    """
    Write a full domain section with findings, evidence chains, and blind spots.

    This matches the Horizon report format:
      - Grade + pillar header
      - Document count + confidence + finding count
      - Domain summary narrative
      - Per-finding: severity, title, description, signals, document quotes,
        business impact, ask-the-target question
      - Blind spots with ⚠ markers
    """
    from docx.shared import Pt, RGBColor

    plabel = dinfo.get("pillar_label", PILLAR_LABELS.get(pid, pid))
    grade = dinfo.get("grade", "UNKNOWN")
    confidence = dinfo.get("confidence", "—")
    findings = dinfo.get("findings", [])
    doc_count = dinfo.get("documents_analyzed", 0)
    pillar_signals = [s for s in all_signals if _get_pillar_id(s) == pid]

    # Pillar header with grade
    doc.add_heading(f"{grade}  {plabel}", level=2)

    # Metadata line
    meta = f"Documents analyzed: {doc_count}  |  Confidence: {confidence}%  |  Findings: {len(findings)}"
    p = doc.add_paragraph()
    run = p.add_run(meta)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 116, 139)

    # Domain summary
    summary = dinfo.get("domain_summary", "")
    if summary:
        doc.add_paragraph(summary)

    # Findings with full evidence chains
    for finding in sorted(findings, key=lambda f: _sev_order(f.get("severity", "MEDIUM"))):
        _write_finding_detail(doc, finding)

    # Blind spots
    blind_spots = dinfo.get("blind_spots", [])
    if blind_spots:
        for bs in blind_spots:
            bs_text = bs if isinstance(bs, str) else bs.get("description", str(bs))
            p = doc.add_paragraph()
            run = p.add_run(f"\u26A0 {bs_text}")  # ⚠ marker
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 116, 139)


def _write_finding_detail(doc, finding: dict):
    """
    Write a single finding with full evidence traceability.

    Format matches Horizon report:
      [SEVERITY] Finding Title
      Description narrative
      Signal: ... (one per evidence chain)
      Document: filename — "quote"
      Business impact: ...
      Ask the target: ...
    """
    from docx.shared import Pt, RGBColor

    sev = finding.get("severity", "MEDIUM")
    title = finding.get("title", "Untitled")
    desc = finding.get("description", "")
    evidence = finding.get("evidence", finding.get("evidence_chain", []))
    business_impact = finding.get("business_impact", "")
    ask_target = finding.get("ask_target", finding.get("question_for_target", ""))

    # Severity + Title
    p = doc.add_paragraph()
    sev_run = p.add_run(f"[{sev}] ")
    sev_run.bold = True
    sev_run.font.color.rgb = _sev_color(sev)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(11)

    # Description
    if desc:
        doc.add_paragraph(desc)

    # Evidence chains
    if isinstance(evidence, list):
        for ev in evidence:
            if isinstance(ev, dict):
                signal_text = ev.get("signal", ev.get("observation", ""))
                doc_name = ev.get("source_doc", ev.get("document", ""))
                quote = ev.get("evidence_quote", ev.get("quote", ""))

                if signal_text:
                    p = doc.add_paragraph()
                    run = p.add_run(f"Signal: ")
                    run.bold = True
                    run.font.size = Pt(9)
                    p.add_run(signal_text).font.size = Pt(9)

                if doc_name:
                    p = doc.add_paragraph()
                    run = p.add_run(f"Document: ")
                    run.bold = True
                    run.font.size = Pt(9)
                    doc_text = doc_name
                    if quote:
                        doc_text += f' — "{quote}"'
                    p.add_run(doc_text).font.size = Pt(9)
            elif isinstance(ev, str):
                p = doc.add_paragraph()
                run = p.add_run(f"Signal: ")
                run.bold = True
                run.font.size = Pt(9)
                p.add_run(ev).font.size = Pt(9)

    # Business impact
    if business_impact:
        p = doc.add_paragraph()
        run = p.add_run("Business impact: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(71, 85, 105)
        p.add_run(business_impact).font.size = Pt(9)

    # Ask the target
    if ask_target:
        p = doc.add_paragraph()
        run = p.add_run("Ask the target: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 64, 175)
        p.add_run(ask_target).font.size = Pt(9)

    doc.add_paragraph("")  # spacer


def _write_signal_detail(doc, sig: dict):
    """
    Write a single VDR signal with evidence (for VDR-only reports without agent findings).

    Produces a compact format:
      [RATING] Title
      Observation text
      Source: document name
      Evidence: "quote"
    """
    from docx.shared import Pt, RGBColor

    rating = sig.get("rating", "UNKNOWN").upper()
    title = sig.get("title", sig.get("signal_id", "Untitled"))
    observation = sig.get("observation", "")
    source = sig.get("source_doc", "")
    evidence = sig.get("evidence_quote", "")
    confidence = sig.get("confidence", "")
    deal_impact = sig.get("deal_implication", "")
    catalog_id = sig.get("catalog_signal_id", sig.get("signal_id", ""))

    # Rating + Title
    p = doc.add_paragraph()
    rating_run = p.add_run(f"[{rating}] ")
    rating_run.bold = True
    rating_run.font.color.rgb = _rating_color(rating)
    title_run = p.add_run(title)
    title_run.bold = True

    # Observation
    if observation:
        doc.add_paragraph(observation)

    # Signal ID
    if catalog_id:
        p = doc.add_paragraph()
        run = p.add_run(f"Signal: ")
        run.bold = True
        run.font.size = Pt(9)
        p.add_run(catalog_id).font.size = Pt(9)

    # Source document + quote
    if source:
        p = doc.add_paragraph()
        run = p.add_run(f"Document: ")
        run.bold = True
        run.font.size = Pt(9)
        doc_text = source
        if evidence:
            doc_text += f' — "{evidence[:300]}"'
        p.add_run(doc_text).font.size = Pt(9)

    # Deal implication
    if deal_impact:
        p = doc.add_paragraph()
        run = p.add_run("Business impact: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(71, 85, 105)
        p.add_run(deal_impact).font.size = Pt(9)

    # Confidence
    if confidence:
        p = doc.add_paragraph()
        run = p.add_run(f"Confidence: {confidence}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(148, 163, 184)

    doc.add_paragraph("")  # spacer


def _build_chase_from_signals(signals: list) -> list:
    """
    Build chase questions from VDR signals (when no agent chase list exists).

    Extracts deal_implication from RED and YELLOW signals to form questions.
    """
    questions = []
    for sig in signals:
        rating = sig.get("rating", "").upper()
        if rating not in ("RED", "YELLOW"):
            continue
        impl = sig.get("deal_implication", "")
        if not impl:
            continue
        pid = _get_pillar_id(sig)
        plabel = PILLAR_LABELS.get(pid, pid)
        questions.append({
            "question": impl,
            "pillar_label": plabel,
            "pillar_id": pid,
            "priority": "high" if rating == "RED" else "medium",
            "source_finding": sig.get("title", sig.get("signal_id", "")),
        })
    return questions


def _sev_order(severity: str) -> int:
    """Sort order for severity (0 = most critical)."""
    return {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}.get(severity, 99)


def _sev_color(severity: str):
    """Return RGBColor for severity."""
    from docx.shared import RGBColor
    return {
        "CRITICAL": RGBColor(220, 38, 38),
        "HIGH": RGBColor(234, 88, 12),
        "MEDIUM": RGBColor(217, 119, 6),
        "LOW": RGBColor(22, 163, 74),
    }.get(severity, RGBColor(107, 114, 128))


def _rating_color(rating: str):
    """Return RGBColor for signal rating."""
    from docx.shared import RGBColor
    return {
        "RED": RGBColor(220, 38, 38),
        "YELLOW": RGBColor(217, 119, 6),
        "GREEN": RGBColor(22, 163, 74),
    }.get(rating, RGBColor(107, 114, 128))
